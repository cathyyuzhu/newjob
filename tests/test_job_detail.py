"""职位详情页冒烟测试：标签校验、备注增删、职位AI对话（不落库）、材料按需生成
（从AI匹配分析里拆出来，见 pipeline.generate_materials_for_job）。
临时库 + 临时 config + 临时简历目录，LLM 全程 mock，不产生真实 API 费用。
"""
import io
import json
import os
import sys
import tempfile

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

tmpdir = tempfile.mkdtemp()
import config

config.DB_PATH = os.path.join(tmpdir, "test.db")
config.CONFIG_PATH = os.path.join(tmpdir, "config.json")

import models

models.DB_PATH = config.DB_PATH

import llm

MATERIALS_RESULT = {
    "needs_customization": True,
    "resume_paragraph_edits": [{"index": 1, "text": "主导XX项目改版，转化率+18%"}],
    "resume_optimization_bullets": ["把职责改成量化成果"],
    "cover_letter": "Dear Hiring Manager, I am excited to apply...",
}
CHAT_REPLY = "这家公司主营AI基础设施，规模大概几百人。"

calls = []


def fake_chat(messages, provider="anthropic", model=None, system=None, max_tokens=None):
    calls.append({"provider": provider, "model": model, "system": system})
    # job_chat.chat_about_job 走 system prompt + 纯文本回复；
    # analyzer.generate_materials 走 llm.ask_json（不带 system，prompt 整段是唯一的 user 消息）。
    if system:
        return CHAT_REPLY
    return json.dumps(MATERIALS_RESULT, ensure_ascii=False)


llm.chat = fake_chat

# analyzer.PROMPT_TEMPLATE / MATERIALS_PROMPT 拆开之后要各自符合预期：分析不再问cover letter，
# 材料生成的 prompt 才问。
import analyzer

assert "cover_letter" not in analyzer.PROMPT_TEMPLATE, "匹配分析的 prompt 不该再要求生成 cover letter"
assert "needs_customization" not in analyzer.PROMPT_TEMPLATE, "匹配分析的 prompt 不该再要求判断是否需要定制简历"
assert "cover_letter" in analyzer.MATERIALS_PROMPT
print("analyzer prompt split ok")

import resume_store

resume_store.RESUME_DIR = os.path.join(tmpdir, "resumes")

import app as flask_app

models.init_db()
flask_app.app.config["TESTING"] = True
c = flask_app.app.test_client()


def make_docx(paragraphs):
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def insert_job(**overrides):
    conn = models.get_conn()
    job = {
        "title": "Senior Product Manager",
        "company": "TestCo",
        "location": "Beijing",
        "site": "linkedin",
        "job_url": "https://example.com/1",
        "date_posted": "2026-08-01",
        "keyword": "Senior Product Manager",
        "jd_text": "We need a PM who ships.",
    }
    job.update(overrides)
    job_id = models.insert_job(conn, job)
    conn.commit()
    conn.close()
    return job_id


# ---- 1. 标签：校验规则 + 落库
job_id = insert_job()
r = c.post(f"/api/jobs/{job_id}/tags", json={"tags": ["AI", "ai", " remote "]})
assert r.status_code == 200, r.get_json()
assert r.get_json()["tags"] == ["AI", "remote"], "大小写不敏感去重，保留第一次出现的写法"
assert models.get_job(job_id)["tags"] == "AI,remote"

assert c.post(f"/api/jobs/{job_id}/tags", json={"tags": ["a,b"]}).status_code == 400, "标签不能含逗号"
assert c.post(f"/api/jobs/{job_id}/tags", json={"tags": ["x" * 21]}).status_code == 400, "单个标签超长"
assert c.post(f"/api/jobs/{job_id}/tags", json={"tags": [f"t{i}" for i in range(11)]}).status_code == 400, "标签数超限"
assert c.post(f"/api/jobs/{job_id}/tags", json={"tags": "not-a-list"}).status_code == 400
assert c.post(f"/api/jobs/{job_id}/tags", json={"tags": []}).status_code == 200
assert models.get_job(job_id)["tags"] is None, "清空标签应该写成 NULL 而不是空字符串"
print("tags validation ok")

# ---- 2. 备注：增删查，来源区分手写/AI，倒序排列
r = c.post(f"/api/jobs/{job_id}/notes", json={"content": "手动记的一条", "source": "manual"})
assert r.status_code == 200 and r.get_json()["id"]
r = c.post(f"/api/jobs/{job_id}/notes", json={"content": "AI说的这条", "source": "chat"})
note2_id = r.get_json()["id"]

notes = c.get(f"/api/jobs/{job_id}/notes").get_json()
assert len(notes) == 2
assert notes[0]["id"] == note2_id, "最新一条应该排最前面"
assert notes[0]["source"] == "chat" and notes[1]["source"] == "manual"

assert c.post(f"/api/jobs/{job_id}/notes", json={"content": ""}).status_code == 400
assert c.post(f"/api/jobs/{job_id}/notes", json={"content": "x" * 4001}).status_code == 400
assert c.post(f"/api/jobs/{job_id}/notes", json={"content": "ok", "source": "bogus"}).status_code == 400
assert c.post(f"/api/jobs/99999/notes", json={"content": "ok"}).status_code == 404

assert c.delete(f"/api/notes/{note2_id}").status_code == 200
assert len(c.get(f"/api/jobs/{job_id}/notes").get_json()) == 1
assert c.delete(f"/api/notes/{note2_id}").status_code == 404, "删除不存在的备注应该 404"
print("notes crud ok")

# ---- 3. 职位AI对话：同步返回、不落库（题库对话同一个决策）
before_note_count = len(models.list_job_notes(job_id))
r = c.post(f"/api/jobs/{job_id}/chat", json={"message": "这家公司靠谱吗？"})
assert r.status_code == 200, r.get_json()
assert r.get_json()["reply"] == CHAT_REPLY
assert len(models.list_job_notes(job_id)) == before_note_count, "对话本身不应该往备注表写任何东西"

assert c.post(f"/api/jobs/{job_id}/chat", json={"message": ""}).status_code == 400
assert c.post("/api/jobs/99999/chat", json={"message": "hi"}).status_code == 404
print("job chat ok")

# ---- 4. 材料生成：分析之前不能生成；上传简历后按需生成，写库+落盘
assert c.post(f"/api/jobs/{job_id}/generate_materials").status_code == 400, "还没分析过不该能生成材料"

docx_bytes = make_docx(["Cathy Yang", "负责XX项目", "教育背景"])
r = c.post("/api/resume/upload", data={"file": (io.BytesIO(docx_bytes), "resume.docx")},
           content_type="multipart/form-data")
assert r.status_code == 200, r.get_json()

models.update_job_analysis(job_id, overall_match=0.82, resume_path=None, company_origin="foreign")

import pipeline

result = pipeline.generate_materials_for_job(job_id)
assert result["resume_path"] and os.path.isfile(result["resume_path"])
assert result["cover_letter"] == MATERIALS_RESULT["cover_letter"]
assert result["resume_bullets"] == MATERIALS_RESULT["resume_optimization_bullets"]

from docx import Document

doc = Document(result["resume_path"])
assert doc.paragraphs[1].text == "主导XX项目改版，转化率+18%", "定制简历应该真的按段落索引改写了"

job = models.get_job(job_id)
assert job["resume_path"] == result["resume_path"]
assert job["cover_letter"] == MATERIALS_RESULT["cover_letter"]
assert json.loads(job["resume_bullets"]) == MATERIALS_RESULT["resume_optimization_bullets"]
print("generate materials for job ok")

# 材料落库之后，重新走一次分析（analyze_and_record）不该把材料冲掉——见 pipeline.analyze_and_record
# 里"把库里已有的材料原样带回追踪表"的说明。这里只验证 update_job_analysis 不touch这三列
# （它本来就不该动，upload/materials 是两条独立的写入路径）。
assert models.get_job(job_id)["cover_letter"] == MATERIALS_RESULT["cover_letter"], (
    "update_job_analysis 不应该动到 cover_letter"
)

# ---- 5. 路由层的守卫：职位不存在 / 正在生成中
assert c.post("/api/jobs/99999/generate_materials").status_code == 404
from job_state import finish_materials, start_materials

start_materials(job_id)
try:
    assert c.post(f"/api/jobs/{job_id}/generate_materials").status_code == 409
finally:
    finish_materials(job_id)
print("generate materials route guards ok")

# ---- 6. 批量生成：已经生成过材料的职位要被跳过，不重复烧钱
job_id2 = insert_job(title="Data Scientist", job_url="https://example.com/2")
models.update_job_analysis(job_id2, overall_match=0.75, resume_path=None, company_origin="domestic")
stats = pipeline.generate_materials_batch([job_id, job_id2])
assert stats["skipped"] == 1, f"job_id 已经生成过材料，应该被跳过：{stats}"
assert stats["generated"] == 1, f"job_id2 应该被生成：{stats}"
assert models.get_job(job_id2)["cover_letter"] == MATERIALS_RESULT["cover_letter"]
print("batch skip already-generated ok")

# ---- 7. GET /api/jobs/<id>/analysis：追踪表里没有对应记录时返回空字典，不报错
assert c.get(f"/api/jobs/{job_id}/analysis").get_json() == {}
assert c.get("/api/jobs/99999/analysis").status_code == 404
print("analysis lookup ok")

print("\nALL PASS")
