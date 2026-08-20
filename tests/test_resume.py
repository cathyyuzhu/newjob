"""「我的简历」冒烟测试：上传校验、体检落库、优化版生成、没简历时的 need_resume 引导。
临时库 + 临时 config + 临时 resumes 目录，LLM 全程 mock，不产生真实 API 费用。
"""
import io
import json
import os
import sys
import tempfile
import time

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
sys.path.insert(0, BASE)

tmpdir = tempfile.mkdtemp()
import config

config.DB_PATH = os.path.join(tmpdir, "test.db")
config.CONFIG_PATH = os.path.join(tmpdir, "config.json")

import models

models.DB_PATH = config.DB_PATH

import llm

FAKE_REVIEW = {
    "overall_score": 0.62,
    "dimension_scores": {"structure": 0.7, "impact": 0.45, "keyword": 72, "clarity": 0.6},
    "summary": "结构清楚，但通篇在写职责而不是成果。",
    "strengths": ["项目背景交代得很干净"],
    "issues": [
        {"severity": "low", "title": "小问题", "detail": "措辞可以更利落", "paragraph_index": 2},
        {"severity": "high", "title": "缺少量化成果", "detail": "第1段只说做了什么", "paragraph_index": 1},
        {"severity": "bogus", "title": "严重度写错了", "detail": "应该被归一化成 medium", "paragraph_index": None},
    ],
    "keyword_coverage": {"covered": ["B端产品"], "missing": ["数据分析", ""]},
    "paragraph_edits": [
        {"index": 1, "original": "负责XX项目", "text": "主导XX项目改版，转化率+18%", "reason": "换成成果导向"},
        {"index": 99, "original": "不存在的段落", "text": "越界的索引，写库时会被跳过", "reason": "-"},
        {"index": 2, "original": "有原文", "text": "   ", "reason": "改写是空的，应该被丢掉"},
    ],
}

calls = []


def fake_chat(messages, provider="anthropic", model=None, system=None, max_tokens=None):
    calls.append({"provider": provider, "model": model, "prompt": messages[0]["content"]})
    # 体检现在是后台线程跑的（见 review_resume_route），留一点延迟让"正在跑的时候
    # 再点一次应该 409"这条断言不会因为跑得太快而变成偶发失败
    time.sleep(0.05)
    return "```json\n" + json.dumps(FAKE_REVIEW, ensure_ascii=False) + "\n```"


llm.chat = fake_chat

import resume_store

# 上传的文件落到临时目录，别往真实的项目 resumes/ 里写
resume_store.RESUME_DIR = os.path.join(tmpdir, "resumes")

import job_state
import app as flask_app


def wait_resume_review(timeout=5.0):
    """体检在后台线程里跑（见 app.py review_resume_route），POST 立刻返回，结果要等
    job_state.resume_review_generating() 翻回 False 才算跑完，跟 test_bank.py 里等
    job_state.bank_generating() 的套路一样。"""
    deadline = time.time() + timeout
    while job_state.resume_review_generating():
        assert time.time() < deadline, "体检后台线程超时没跑完"
        time.sleep(0.05)

models.init_db()
flask_app.app.config["TESTING"] = True
c = flask_app.app.test_client()


def make_docx(paragraphs):
    """造一份真的 docx（不 mock read_resume_text）：上传校验的关键一步就是"能不能真的
    被 python-docx 解析出正文"，用假的就把这条测掉了。"""
    from docx import Document

    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---- 1. 还没上传简历：所有依赖简历的入口都要回 409 + need_resume，而不是笼统的 500
assert c.get("/api/resume").get_json() == {"exists": False}

_conn = models.get_conn()
job_id = models.insert_job(
    _conn,
    {
        "title": "Senior Product Manager",
        "company": "TestCo",
        "location": "Beijing",
        "site": "linkedin",
        "job_url": "https://example.com/1",
        "date_posted": "2026-08-01",
        "keyword": "Senior Product Manager",
        "jd_text": "We need a PM who ships.",
    },
)
_conn.commit()
_conn.close()

for path in [f"/api/jobs/{job_id}/analyze", "/api/jobs/analyze_all",
             f"/api/jobs/{job_id}/interview_prep", "/api/interview/bank/generate"]:
    r = c.post(path)
    assert r.status_code == 409, f"{path} 应该回 409，实际 {r.status_code}"
    assert r.get_json().get("need_resume") is True, f"{path} 缺少 need_resume 标记"

r = c.post("/api/resume/review")
assert r.status_code == 409 and r.get_json()["need_resume"] is True
# 没简历也不该把职位标成"分析失败"——那不是这条职位的问题
assert not models.get_job(job_id).get("analysis_error"), "没上传简历不该给职位盖上分析失败"
print("need_resume guards ok")

# ---- 2. 上传校验：扩展名、假 docx
r = c.post("/api/resume/upload", data={"file": (io.BytesIO(b"%PDF-1.4 fake"), "resume.pdf")},
           content_type="multipart/form-data")
assert r.status_code == 400 and ".docx" in r.get_json()["error"]

# 把 PDF 改个后缀传上来：前两关（扩展名/大小）都过得了，必须靠"试着解析"这一关拦下
r = c.post("/api/resume/upload", data={"file": (io.BytesIO(b"%PDF-1.4 fake"), "resume.docx")},
           content_type="multipart/form-data")
assert r.status_code == 400, r.get_json()
assert not resume_store.has_base_resume(), "解析失败的文件不该被记成当前简历"
leftovers = os.listdir(resume_store.RESUME_DIR) if os.path.isdir(resume_store.RESUME_DIR) else []
assert not leftovers, f"解析失败后不该留下残留文件：{leftovers}"
print("upload validation ok")

# ---- 3. 上传合法 docx
docx_bytes = make_docx(["Cathy Yang", "负责XX项目", "有原文", "教育背景"])
r = c.post("/api/resume/upload", data={"file": (io.BytesIO(docx_bytes), "我的简历 v3.docx")},
           content_type="multipart/form-data")
assert r.status_code == 200, r.get_json()
meta = r.get_json()
assert meta["exists"] and meta["paragraph_count"] == 4, meta
# 原始文件名（含中文和空格）要原样留着展示，磁盘名则是时间戳
assert meta["filename"] == "我的简历 v3.docx", meta
assert os.path.basename(config.load_config()["base_resume_path"]).startswith("base_")
assert config.load_config()["base_resume_meta"]["original_filename"] == "我的简历 v3.docx"
assert c.get("/api/resume/download").status_code == 200
print("upload ok")

# ---- 4. 体检：LLM 的脏数据要被归一化，坏的 paragraph_edits 要被丢掉
# 后台线程跑，POST 立刻返回 {"started": True}，不再直接带结果——等 generating 标志
# 翻回 False（wait_resume_review），结果从 GET /api/resume/review 里读
r = c.post("/api/resume/review")
assert r.status_code == 200 and r.get_json()["started"] is True, r.get_json()
# 正在跑的时候再点一次应该 409，不是重复触发一次新的 LLM 调用
assert c.post("/api/resume/review").status_code == 409
wait_resume_review()
row = c.get("/api/resume/review").get_json()
assert row["generating"] is False and not row["background_error"], row
content = json.loads(row["content_json"])
# 72 是百分制写法，要收成 0.72
assert abs(content["dimension_scores"]["keyword"] - 0.72) < 1e-6, content["dimension_scores"]
assert 0 <= content["overall_score"] <= 1
# 严重度非法的归 medium，且整体按 high → medium → low 排好
assert [i["severity"] for i in content["issues"]] == ["high", "medium", "low"], content["issues"]
assert content["keyword_coverage"]["missing"] == ["数据分析"], "空字符串应该被过滤掉"
# 三条 edits 里，越界索引和空改写各丢一条，只剩能真正喂给 write_tailored_resume 的那条
assert len(content["paragraph_edits"]) == 1, content["paragraph_edits"]
assert content["paragraph_edits"][0]["index"] == 1

# 体检用的是 resume_review 功能位，而且 prompt 里带上了配置里的目标岗位方向
assert calls, "没有调用到 LLM"
assert "Senior Product Manager" in calls[-1]["prompt"], "体检 prompt 里应该有目标岗位方向"
assert row["id"] and row["stale"] is False, row
print("resume review ok")

# ---- 5. 换一份简历之后，旧体检结果要被标成过期（段落索引已经对不上了）
r = c.post("/api/resume/upload",
           data={"file": (io.BytesIO(make_docx(["Cathy Yang", "改过的第二段", "第三段"])), "v4.docx")},
           content_type="multipart/form-data")
assert r.status_code == 200
assert c.get("/api/resume/review").get_json()["stale"] is True, "换简历后旧体检该标记为过期"
print("stale review flag ok")

# ---- 6. 优化版 docx：勾中的段落被真的改写，没勾的原样不动
r = c.post("/api/resume/optimize", json={"edits": [{"index": 1, "text": "主导XX项目改版，转化率+18%"}]})
assert r.status_code == 200, r.get_json()
assert r.get_json()["download_name"] == "v4_优化版.docx", r.get_json()

from docx import Document

doc = Document(resume_store.optimized_path())
texts = [p.text for p in doc.paragraphs]
assert texts[1] == "主导XX项目改版，转化率+18%", texts
assert texts[0] == "Cathy Yang" and texts[2] == "第三段", "没勾的段落不该被动"
assert c.get("/api/resume/optimized").status_code == 200

# 一条都没勾就点生成，要给一句人话而不是 500
assert c.post("/api/resume/optimize", json={"edits": []}).status_code == 400
print("optimized docx ok")

# ---- 7. 定制简历列表：文件没了的要标出来，而不是给个点了才 404 的下载按钮
models.update_job_analysis(job_id, overall_match=0.81,
                           resume_path=os.path.join(tmpdir, "gone.docx"), company_origin="foreign")
rows = c.get("/api/resume/tailored").get_json()
assert len(rows) == 1 and rows[0]["file_exists"] is False, rows
print("tailored list ok")

# ---- 8. 删除：清引用 + 删文件，之后又回到"要先上传"的状态
assert c.delete("/api/resume").get_json() == {"exists": False}
assert config.load_config()["base_resume_path"] == ""
assert c.post(f"/api/jobs/{job_id}/analyze").status_code == 409
print("delete ok")

print("\nALL PASS")
