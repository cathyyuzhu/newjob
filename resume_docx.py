"""Reads the user's base resume docx as plain text, and can produce a
tailored copy by replacing the text of specific paragraphs (identified by
index) while keeping the original document's styling/formatting intact.
"""
import os

from docx import Document


def read_resume_text(path):
    if not path or not os.path.exists(path):
        return None
    doc = Document(path)
    lines = []
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text:
            lines.append(f"[{i}] {text}")
    return "\n".join(lines)


def write_tailored_resume(base_path, output_path, paragraph_edits):
    """paragraph_edits: list of {"index": int, "text": str} — replaces the
    text of doc.paragraphs[index] with `text`, keeping the first run's
    formatting (bold/font/etc.) and dropping any extra runs in that
    paragraph so the replacement text renders cleanly.
    """
    doc = Document(base_path)
    paragraphs = doc.paragraphs
    for edit in paragraph_edits:
        idx = edit.get("index")
        text = edit.get("text", "")
        if idx is None or idx < 0 or idx >= len(paragraphs):
            continue
        p = paragraphs[idx]
        if not p.runs:
            p.add_run(text)
            continue
        p.runs[0].text = text
        for extra_run in p.runs[1:]:
            extra_run.text = ""
    doc.save(output_path)
    return output_path
